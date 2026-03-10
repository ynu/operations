#!/usr/bin/env python3
"""
将 Word 文档 (.docx) 转换为 Markdown 格式
使用方式: python convert_docx_to_md.py <input.docx> <output.md>
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys
import os
import re

def get_text_style(run):
    """获取文本的样式标记"""
    styles = []
    if run.bold:
        styles.append('bold')
    if run.italic:
        styles.append('italic')
    if run.underline:
        styles.append('underline')
    if run.font.strike:
        styles.append('strike')
    return styles

def apply_text_style(text, styles):
    """应用 Markdown 样式到文本"""
    if not text:
        return text
    
    # 处理粗体
    if 'bold' in styles:
        text = f"**{text}**"
    # 处理斜体
    if 'italic' in styles:
        text = f"*{text}*"
    # 处理删除线
    if 'strike' in styles:
        text = f"~~{text}~~"
    
    return text

def convert_table_to_md(table):
    """将 Word 表格转换为 Markdown 表格"""
    if not table.rows:
        return ""
    
    md_lines = []
    
    # 处理表头
    header_row = table.rows[0]
    header_cells = [cell.text.strip() for cell in header_row.cells]
    md_lines.append("| " + " | ".join(header_cells) + " |")
    md_lines.append("|" + "|".join([" --- " for _ in header_cells]) + "|")
    
    # 处理数据行
    for row in table.rows[1:]:
        cells = [cell.text.strip() for cell in row.cells]
        md_lines.append("| " + " | ".join(cells) + " |")
    
    return "\n".join(md_lines)

def convert_paragraph_to_md(paragraph):
    """将 Word 段落转换为 Markdown"""
    text_parts = []
    
    for run in paragraph.runs:
        text = run.text
        if text:
            styles = get_text_style(run)
            styled_text = apply_text_style(text, styles)
            text_parts.append(styled_text)
    
    return "".join(text_parts)

def convert_docx_to_md(docx_path, md_path):
    """将 Word 文档转换为 Markdown"""
    doc = Document(docx_path)
    
    md_content = []
    
    for element in doc.element.body:
        # 处理段落
        if element.tag.endswith('p'):
            paragraph = next((p for p in doc.paragraphs if p._element == element), None)
            if paragraph:
                text = convert_paragraph_to_md(paragraph)
                
                # 跳过空段落
                if not text.strip():
                    md_content.append("")
                    continue
                
                # 根据样式判断标题级别
                style_name = paragraph.style.name if paragraph.style else ""
                
                if style_name.startswith('Heading'):
                    level = style_name.replace('Heading ', '')
                    try:
                        level = int(level)
                        md_content.append(f"{'#' * level} {text}")
                    except ValueError:
                        md_content.append(text)
                elif paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                    md_content.append(f"<center>{text}</center>")
                else:
                    md_content.append(text)
        
        # 处理表格
        elif element.tag.endswith('tbl'):
            table = next((t for t in doc.tables if t._element == element), None)
            if table:
                md_table = convert_table_to_md(table)
                if md_table:
                    md_content.append("")
                    md_content.append(md_table)
                    md_content.append("")
    
    # 写入 Markdown 文件
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write("\n".join(md_content))
    
    print(f"转换完成: {md_path}")

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("用法: python convert_docx_to_md.py <input.docx> <output.md>")
        print("示例: python convert_docx_to_md.py ./api_doc.docx ./output.md")
        sys.exit(1)
    
    docx_file = sys.argv[1]
    md_file = sys.argv[2]
    
    # 检查输入文件是否存在
    if not os.path.exists(docx_file):
        print(f"错误: 输入文件不存在: {docx_file}")
        sys.exit(1)
    
    # 检查是否是 .docx 文件
    if not docx_file.endswith('.docx'):
        print(f"错误: 输入文件必须是 .docx 格式")
        sys.exit(1)
    
    convert_docx_to_md(docx_file, md_file)
